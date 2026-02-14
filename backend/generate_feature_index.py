"""Generate Feature Index Word Document"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Title
title = doc.add_heading('D&V Business Consulting', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Business Management Application — Feature Index')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# Helper
def add_section(title, items, status_map=None):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Feature'
    hdr[2].text = 'Status'
    hdr[3].text = 'Details'
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i, item in enumerate(items, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = item[0]
        row[2].text = item[1]
        row[3].text = item[2]
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph('')

# ===== AUTHENTICATION & USER MANAGEMENT =====
add_section('1. Authentication & User Management', [
    ('JWT Email/Password Authentication', 'DONE', 'Secure login with token-based auth'),
    ('13 Customizable User Roles', 'DONE', 'Admin, Manager, Executive, Consultant, PM, Principal Consultant, Lean/Lead/Senior Consultant, HR Executive/Manager, Account Manager, SME'),
    ('Role & Permissions Management', 'DONE', 'Per-module CRUD + special actions, system vs custom roles'),
    ('User Management Page', 'DONE', 'Users tab + Roles tab, add user, change roles, permissions dialog'),
    ('User Profile Page', 'DONE', 'View/edit profile, role badge, permissions view'),
    ('Consultant-specific Dashboard', 'DONE', 'Separate dashboard for consultant role'),
])

# ===== SALES WORKFLOW =====
add_section('2. Sales Workflow', [
    ('Lead Management', 'DONE', 'Create/edit leads, automated scoring, status tracking, high-priority display'),
    ('Pricing Plans', 'DONE', 'Consultant allocation, duration, discounts, GST calculation'),
    ('SOW Builder', 'DONE', 'Spreadsheet-style inline editing, categories, version tracking, freeze after kick-off'),
    ('SOW Role-Based Segregation', 'DONE', 'Sales team: create/edit | Consulting: update status | PM: approve/reject'),
    ('SOW Inline Editing & Roadmap View', 'DONE', 'Inline row add, consultant assignment, start week, monthly roadmap, Gantt chart'),
    ('SOW Documents Per Item', 'DONE', 'Per-line-item document upload/download, auto-notification on completion'),
    ('SOW Version History', 'DONE', 'Full snapshot per version, before/after change highlighting'),
    ('Quotations', 'DONE', 'Generate from pricing plans, finalize for agreements'),
    ('Agreements', 'DONE', 'Create from quotations, 12-section structure (NDA, NCA, SOW, Pricing, etc.)'),
    ('Agreement & SOW Export (Word/PDF)', 'DONE', 'Professional document generation with company branding, tables, signatures'),
    ('Manager Approvals for Agreements', 'DONE', 'Approve/reject with comments'),
    ('Sales Meetings & MOM', 'DONE', 'Lightweight form, lead linking, simple MOM (agenda, decisions, notes)'),
    ('Client Master Module', 'DONE', 'Company info, contacts (SPOCs), revenue history, industry tracking'),
    ('Sales Reports (8 reports)', 'DONE', 'Lead Summary, Conversion Funnel, Source Analysis, Client Overview, Industry Breakdown, Pipeline, Quotation, Agreement'),
    ('Drag-and-Drop Gantt Chart', 'PENDING', 'Draggable bars for date/week adjustment'),
    ('RACI Matrix for SOW', 'PENDING', 'Inline-editable role assignment, PDF/Excel export'),
    ('Rocket Reach Integration', 'PENDING', 'Lead enrichment via Rocket Reach API'),
])

# ===== CONSULTING & PROJECT MANAGEMENT =====
add_section('3. Consulting & Project Management', [
    ('Project Management', 'DONE', 'Create projects, assign consultants, track status, budget, meetings'),
    ('Project Assignment System', 'DONE', 'Assign/change/unassign consultants, track meetings per assignment'),
    ('Project Roadmap', 'DONE', 'Monthly phased plans, Table + Kanban views, submit to client, inline status updates'),
    ('Consulting Meetings & MOM', 'DONE', 'Project-linked, detailed MOM with action items, send to client, commitment tracking'),
    ('Commitment Tracking', 'DONE', 'Committed vs actual meetings per project with variance & completion %'),
    ('Consultant Management', 'DONE', 'List view, bandwidth tracking (online/offline/mixed limits), project allocation'),
    ('Consultant Performance Metrics', 'DONE', 'Configurable per project by Principal Consultant, admin approval required'),
    ('Default Metrics', 'DONE', 'SOW Delivery (20%), Roadmap Achievement (20%), Records Timeliness (15%), SOW Quality (25%), Meeting Adherence (20%)'),
    ('Performance Scoring', 'DONE', 'RM rates monthly (0-100 per metric), weighted overall score, summary view'),
    ('Task Management', 'DONE', 'CRUD with categories, priorities, status, timeline/Gantt view'),
    ('Handover Alerts', 'DONE', '15-day deadline tracking, color-coded urgency, create project/assign consultants'),
    ('Kick-off Meeting & SOW Freeze', 'DONE', 'Schedule by Principal Consultant, SOW freezes, admin override'),
    ('Action Items & Follow-up Tasks', 'DONE', 'From MOM, assigned with due dates, auto-task creation, manager notification'),
    ('Consulting Reports (5 reports)', 'DONE', 'SOW Status, Project Summary, Consultant Allocation, Approval Turnaround, Pending Approvals'),
    ('SOW Linkages with Performance', 'DONE', 'SOW quality scored by RM, document approval rating as metric'),
    ('Project Stage Tracking', 'PENDING', 'Detailed stage-wise progress tracking'),
])

# ===== HR MODULE =====
add_section('4. HR Module', [
    ('Employee Directory', 'DONE', 'Searchable list, department filters, employee ID, reporting manager'),
    ('Employee CRUD', 'DONE', 'Personal info, work info, financial, leave balance, documents'),
    ('User-Employee Linking', 'DONE', 'Link system users to employee records, sync from users'),
    ('Org Chart (Visual Hierarchy)', 'DONE', 'Expandable tree view, department color coding, manager/subordinate display'),
    ('Leave Management (HR Admin)', 'DONE', 'All requests view, approval tracking, RM → HR Manager flow'),
    ('Attendance Management (HR Admin)', 'DONE', 'Manual entry, CSV bulk upload, monthly summary, daily records'),
    ('Payroll', 'DONE', 'Salary components config, generate slips (individual/bulk), earnings & deductions'),
    ('Salary Components', 'DONE', 'Basic (40%), HRA (20%), Special Allowance (20%), Conveyance (₹1,600), Medical (₹1,250) | PF (12%), PT (₹200), ESI (0.75%)'),
    ('Expense Reimbursement in Payroll', 'DONE', 'Approved expenses auto-included as "Conveyance Reimbursement" in salary slip'),
    ('Expense Request System', 'DONE', 'Draft → Pending → Approved → Reimbursed, line items, client/project linking'),
    ('HR Reports (3 reports)', 'DONE', 'Employee Directory, Department Analysis, Leave Utilization'),
    ('Finance Reports (3 reports)', 'DONE', 'Client Revenue, Expense Summary, Expense Category Analysis'),
    ('Detailed Time Tracking', 'PENDING', 'Hourly time tracking per project/task'),
    ('HR Workflow Module', 'PENDING', 'Onboarding, offboarding, probation workflows'),
])

# ===== SELF-SERVICE (MY WORKSPACE) =====
add_section('5. My Workspace (All Users)', [
    ('My Attendance', 'DONE', 'View own attendance records, monthly summary with stats (present/absent/WFH/leave)'),
    ('My Leaves', 'DONE', 'Leave balance cards (casual/sick/earned), apply leave, track request status'),
    ('My Salary Slips', 'DONE', 'All historical monthly salary statements, detailed slip view (earnings, deductions, net pay)'),
    ('My Expenses', 'DONE', 'Submit expenses, track claim status (draft/pending/approved/reimbursed), submit for approval'),
])

# ===== APPROVAL WORKFLOW =====
add_section('6. Approval Workflow Engine', [
    ('Approvals Center', 'DONE', 'Central hub with 3 tabs: Pending My Action, My Requests, All Approvals'),
    ('Multi-Level Approval Chain', 'DONE', 'Reporting Manager → HR/Admin based on type'),
    ('Approval Types', 'DONE', 'SOW Items, Agreements, Quotations, Leave Requests, Expenses, Client Communications'),
    ('Performance Metrics Approval', 'DONE', 'Admin approves metrics before they populate to users'),
    ('Visual Approval Chain', 'DONE', 'Chain preview with levels and status'),
])

# ===== REPORTING =====
add_section('7. Reports & Analytics', [
    ('19 Core Reports', 'DONE', 'Sales (8), Finance (3), HR (3), Operations (5)'),
    ('Export Formats', 'DONE', 'Excel (.xlsx) and PDF'),
    ('Domain-Specific Filtering', 'DONE', 'Auto-filter via URL params (?category=hr/sales/operations)'),
    ('Role-Based Access', 'DONE', 'Admin: all | Manager: all | HR Manager: HR+Finance | PM: Ops+Client | Executive: Sales'),
    ('Quick Stats Dashboard', 'DONE', 'Leads, Clients, Employees, Projects, Revenue, Pending Approvals'),
    ('Quarterly Activity Reports', 'PENDING', 'Cross-domain quarterly summaries'),
    ('30+ Extended Reports', 'PENDING', 'Expanded report catalog'),
])

# ===== NAVIGATION & UI =====
add_section('8. Navigation & UI/UX', [
    ('Domain-Segmented Sidebar', 'DONE', '5 sections: My Workspace, HR, Sales, Consulting, Admin'),
    ('Role-Based Nav Visibility', 'DONE', 'Each section visible only to relevant roles'),
    ('Sales Flow Branch', 'DONE', 'Visual connector: Leads → Pricing Plans → Quotations → Agreements'),
    ('Sticky Scrollable Sidebar', 'DONE', 'Fixed sidebar with scrollable nav for many items'),
    ('Black & White Theme', 'DONE', 'Professional monochrome design with Shadcn/UI components'),
    ('Indian Rupees (₹)', 'DONE', 'All currency displayed in INR format'),
    ('Company Logo', 'DONE', 'D&V Business Consulting branding in sidebar'),
])

# ===== PENDING / FUTURE =====
add_section('9. Pending & Future Features', [
    ('Drag-and-Drop Gantt Chart', 'P1 - PENDING', 'Draggable bars for date/week adjustment, performance metrics'),
    ('RACI Matrix for SOW', 'P1 - PENDING', 'Inline-editable role assignment with PDF/Excel export'),
    ('Rocket Reach Integration', 'P2 - PENDING', 'Lead enrichment API'),
    ('Email Sending (SMTP)', 'P2 - PENDING', 'Currently MOCKED — requires SMTP credentials'),
    ('Detailed Time Tracking', 'P2 - PENDING', 'Hourly tracking per project/task'),
    ('Project Stage Tracking', 'P2 - PENDING', 'Stage-wise progress tracking'),
    ('Quarterly Activity Reports', 'P2 - PENDING', 'Cross-domain quarterly summaries'),
    ('Marketing Flow Module', 'P3 - BACKLOG', 'Marketing campaign and workflow management'),
    ('Finance & Accounts Module', 'P3 - BACKLOG', 'Invoicing, payments, accounting'),
    ('Salary Slip PDF Download', 'P3 - BACKLOG', 'Direct PDF download for employees'),
    ('Refactor server.py', 'TECH DEBT', 'Split monolithic 7500+ line file into modular routers'),
])

# Summary
doc.add_heading('Feature Summary', level=1)
summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Light Grid Accent 1'
hdr = summary_table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Done'
hdr[2].text = 'Pending'
for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True

summary_data = [
    ('Authentication & Users', '6', '0'),
    ('Sales Workflow', '14', '3'),
    ('Consulting & Projects', '14', '1'),
    ('HR Module', '12', '2'),
    ('My Workspace (Self-Service)', '4', '0'),
    ('Approval Workflow', '5', '0'),
    ('Reports & Analytics', '5', '2'),
    ('Navigation & UI', '7', '0'),
    ('TOTAL', '67', '8+'),
]
for data in summary_data:
    row = summary_table.add_row().cells
    row[0].text = data[0]
    row[1].text = data[1]
    row[2].text = data[2]
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    if data[0] == 'TOTAL':
        for cell in row:
            cell.paragraphs[0].runs[0].font.bold = True

doc.save('/app/uploads/Feature_Index_DVB_Consulting.docx')
print("Document saved: /app/uploads/Feature_Index_DVB_Consulting.docx")

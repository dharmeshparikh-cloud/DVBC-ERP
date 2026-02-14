"""
Document Generator Module
Generates Word (.docx) and PDF documents for Agreements and SOW
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Optional


def format_currency(amount: float) -> str:
    """Format amount as Indian Rupees"""
    if amount >= 10000000:  # 1 Crore
        return f"₹{amount/10000000:.2f} Cr"
    elif amount >= 100000:  # 1 Lakh
        return f"₹{amount/100000:.2f} L"
    else:
        return f"₹{amount:,.2f}"


def format_date(date_val) -> str:
    """Format date for display"""
    if not date_val:
        return "TBD"
    if isinstance(date_val, str):
        try:
            date_val = datetime.fromisoformat(date_val.replace('Z', '+00:00'))
        except ValueError:
            return date_val
    return date_val.strftime("%d %B %Y")


class AgreementDocumentGenerator:
    """Generates Agreement documents in Word and PDF formats"""
    
    COMPANY_NAME = "D&V Business Consulting"
    COMPANY_ADDRESS = "Business Address, City, State - PIN"
    
    def __init__(self, agreement_data: Dict[str, Any]):
        self.data = agreement_data
        
    def generate_word(self) -> BytesIO:
        """Generate Word document for the agreement"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Title
        title = doc.add_heading('CONSULTING AGREEMENT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agreement Number and Date
        agreement_num = self.data.get('agreement_number', 'N/A')
        created_at = self.data.get('created_at', '')
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Agreement No: {agreement_num}")
        run.bold = True
        p.add_run(f"\nDate: {format_date(created_at)}")
        
        doc.add_paragraph()
        
        # Section 1: Party Information
        self._add_section_heading(doc, "1. PARTY INFORMATION")
        
        # First Party
        p = doc.add_paragraph()
        p.add_run("FIRST PARTY (Service Provider):").bold = True
        doc.add_paragraph(f"{self.COMPANY_NAME}")
        doc.add_paragraph(f"{self.COMPANY_ADDRESS}")
        
        doc.add_paragraph()
        
        # Second Party
        p = doc.add_paragraph()
        p.add_run("SECOND PARTY (Client):").bold = True
        
        client = self.data.get('client', {})
        party_name = self.data.get('party_name', '')
        if party_name:
            doc.add_paragraph(party_name)
        else:
            client_name = f"{client.get('name', '')}".strip()
            company = client.get('company', '')
            if client_name:
                doc.add_paragraph(client_name)
            if company:
                doc.add_paragraph(company)
        
        if client.get('email'):
            doc.add_paragraph(f"Email: {client.get('email')}")
        if client.get('phone'):
            doc.add_paragraph(f"Phone: {client.get('phone')}")
        
        doc.add_paragraph()
        
        # Section 2: Agreement Between
        self._add_section_heading(doc, "2. AGREEMENT BETWEEN PARTIES")
        p = doc.add_paragraph()
        company_section = self.data.get('company_section', '')
        if company_section:
            p.add_run(company_section)
        else:
            p.add_run(f"This Agreement is entered into between {self.COMPANY_NAME} (hereinafter referred to as \"Consultant\") and the Second Party (hereinafter referred to as \"Client\") for the provision of consulting services as described herein.")
        
        doc.add_paragraph()
        
        # Section 3: Confidentiality
        if self.data.get('confidentiality_clause'):
            self._add_section_heading(doc, "3. CONFIDENTIALITY")
            doc.add_paragraph(self.data.get('confidentiality_clause'))
            doc.add_paragraph()
        
        # Section 4: NDA
        if self.data.get('nda_clause'):
            self._add_section_heading(doc, "4. NON-DISCLOSURE AGREEMENT (NDA)")
            doc.add_paragraph(self.data.get('nda_clause'))
            doc.add_paragraph()
        
        # Section 5: NCA
        if self.data.get('nca_clause'):
            self._add_section_heading(doc, "5. NON-COMPETE AGREEMENT (NCA)")
            doc.add_paragraph(self.data.get('nca_clause'))
            doc.add_paragraph()
        
        # Section 6: Renewal Terms
        if self.data.get('renewal_clause'):
            self._add_section_heading(doc, "6. RENEWAL TERMS")
            doc.add_paragraph(self.data.get('renewal_clause'))
            doc.add_paragraph()
        
        # Section 7: Conveyance
        if self.data.get('conveyance_clause'):
            self._add_section_heading(doc, "7. CONVEYANCE")
            doc.add_paragraph(self.data.get('conveyance_clause'))
            doc.add_paragraph()
        
        # Section 8: Scope of Work
        self._add_section_heading(doc, "8. SCOPE OF WORK (SOW)")
        
        sow_table_data = self.data.get('sow_table', [])
        if sow_table_data:
            # Create SOW table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['#', 'Category', 'Title', 'Description', 'Timeline']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for idx, item in enumerate(sow_table_data, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = item.get('category', '')
                row_cells[2].text = item.get('title', '')
                
                # Description with deliverables
                desc = item.get('description', '')
                deliverables = item.get('deliverables', [])
                if deliverables:
                    desc += "\nDeliverables:\n" + "\n".join([f"• {d}" for d in deliverables])
                row_cells[3].text = desc
                
                timeline = item.get('timeline_weeks')
                row_cells[4].text = f"{timeline} weeks" if timeline else "TBD"
        else:
            doc.add_paragraph("(SOW details to be defined)")
        
        doc.add_paragraph()
        
        # Section 9: Project Details
        self._add_section_heading(doc, "9. PROJECT DETAILS")
        
        project = self.data.get('project_details', {})
        p = doc.add_paragraph()
        p.add_run("Start Date: ").bold = True
        p.add_run(format_date(project.get('start_date')))
        
        p = doc.add_paragraph()
        p.add_run("Duration: ").bold = True
        duration = project.get('duration_months')
        p.add_run(f"{duration} months" if duration else "TBD")
        
        if project.get('payment_schedule'):
            p = doc.add_paragraph()
            p.add_run("Payment Schedule: ").bold = True
            p.add_run(project.get('payment_schedule'))
        
        doc.add_paragraph()
        
        # Section 10: Team Engagement
        self._add_section_heading(doc, "10. TEAM ENGAGEMENT")
        
        team_data = self.data.get('team_engagement', [])
        if team_data and isinstance(team_data, list):
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            headers = ['Consultant Type', 'Count', 'Meetings', 'Hours', 'Rate/Meeting']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for item in team_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('type', '')
                row_cells[1].text = str(item.get('count', 1))
                row_cells[2].text = str(item.get('meetings', 0))
                row_cells[3].text = str(item.get('hours', 0))
                row_cells[4].text = format_currency(item.get('rate', 12500))
        else:
            doc.add_paragraph("(Team details as per SOW)")
        
        doc.add_paragraph()
        
        # Section 11: Pricing
        self._add_section_heading(doc, "11. PRICING PLAN")
        
        pricing = self.data.get('pricing', {})
        
        pricing_table = doc.add_table(rows=5, cols=2)
        pricing_table.style = 'Table Grid'
        
        pricing_items = [
            ("Total Meetings", str(pricing.get('total_meetings', 0))),
            ("Subtotal", format_currency(pricing.get('subtotal', 0))),
            ("Discount", format_currency(pricing.get('discount', 0))),
            ("GST (18%)", format_currency(pricing.get('gst', 0))),
            ("Grand Total", format_currency(pricing.get('total', 0)))
        ]
        
        for i, (label, value) in enumerate(pricing_items):
            pricing_table.rows[i].cells[0].text = label
            pricing_table.rows[i].cells[1].text = value
            if i == 4:  # Grand Total row
                pricing_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                pricing_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Section 12: Payment Terms
        self._add_section_heading(doc, "12. PAYMENT TERMS & CONDITIONS")
        
        p = doc.add_paragraph()
        p.add_run("Payment Terms: ").bold = True
        p.add_run(self.data.get('payment_terms', 'Net 30 days'))
        
        if self.data.get('payment_conditions'):
            doc.add_paragraph(self.data.get('payment_conditions'))
        
        if self.data.get('special_conditions'):
            p = doc.add_paragraph()
            p.add_run("Special Conditions: ").bold = True
            doc.add_paragraph(self.data.get('special_conditions'))
        
        doc.add_paragraph()
        
        # Section 13: Signatures
        self._add_section_heading(doc, "13. SIGNATURES")
        
        doc.add_paragraph()
        
        # Signature table
        sig_table = doc.add_table(rows=4, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        sig_table.rows[0].cells[0].text = "For D&V Business Consulting"
        sig_table.rows[0].cells[1].text = "For Client"
        sig_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        sig_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
        
        # Signature lines
        sig_table.rows[1].cells[0].text = "\n\n_____________________"
        sig_table.rows[1].cells[1].text = "\n\n_____________________"
        
        # Name
        sig_table.rows[2].cells[0].text = "Authorized Signatory"
        sig_table.rows[2].cells[1].text = party_name or "Client Representative"
        
        # Date
        sig_table.rows[3].cells[0].text = "Date: _______________"
        sig_table.rows[3].cells[1].text = "Date: _______________"
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_section_heading(self, doc, title: str):
        """Add a formatted section heading"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
    
    def generate_pdf(self) -> BytesIO:
        """Generate PDF document for the agreement"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20,
            textColor=colors.black
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.black,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        center_style = ParagraphStyle(
            'CustomCenter',
            parent=normal_style,
            alignment=TA_CENTER
        )
        
        elements = []
        
        # Title
        elements.append(Paragraph("CONSULTING AGREEMENT", title_style))
        
        # Agreement info
        agreement_num = self.data.get('agreement_number', 'N/A')
        created_at = self.data.get('created_at', '')
        elements.append(Paragraph(f"<b>Agreement No:</b> {agreement_num}", center_style))
        elements.append(Paragraph(f"Date: {format_date(created_at)}", center_style))
        elements.append(Spacer(1, 20))
        
        # Section 1: Party Information
        elements.append(Paragraph("1. PARTY INFORMATION", heading_style))
        elements.append(Paragraph("<b>FIRST PARTY (Service Provider):</b>", normal_style))
        elements.append(Paragraph(self.COMPANY_NAME, normal_style))
        elements.append(Paragraph(self.COMPANY_ADDRESS, normal_style))
        elements.append(Spacer(1, 10))
        
        elements.append(Paragraph("<b>SECOND PARTY (Client):</b>", normal_style))
        client = self.data.get('client', {})
        party_name = self.data.get('party_name', '')
        if party_name:
            elements.append(Paragraph(party_name, normal_style))
        else:
            if client.get('name'):
                elements.append(Paragraph(client.get('name'), normal_style))
            if client.get('company'):
                elements.append(Paragraph(client.get('company'), normal_style))
        if client.get('email'):
            elements.append(Paragraph(f"Email: {client.get('email')}", normal_style))
        if client.get('phone'):
            elements.append(Paragraph(f"Phone: {client.get('phone')}", normal_style))
        elements.append(Spacer(1, 10))
        
        # Section 2: Agreement Between
        elements.append(Paragraph("2. AGREEMENT BETWEEN PARTIES", heading_style))
        company_section = self.data.get('company_section', '')
        if company_section:
            elements.append(Paragraph(company_section, normal_style))
        else:
            elements.append(Paragraph(f"This Agreement is entered into between {self.COMPANY_NAME} (hereinafter referred to as \"Consultant\") and the Second Party (hereinafter referred to as \"Client\") for the provision of consulting services as described herein.", normal_style))
        
        # Conditional sections
        section_num = 3
        
        if self.data.get('confidentiality_clause'):
            elements.append(Paragraph(f"{section_num}. CONFIDENTIALITY", heading_style))
            elements.append(Paragraph(self.data.get('confidentiality_clause'), normal_style))
            section_num += 1
        
        if self.data.get('nda_clause'):
            elements.append(Paragraph(f"{section_num}. NON-DISCLOSURE AGREEMENT (NDA)", heading_style))
            elements.append(Paragraph(self.data.get('nda_clause'), normal_style))
            section_num += 1
        
        if self.data.get('nca_clause'):
            elements.append(Paragraph(f"{section_num}. NON-COMPETE AGREEMENT (NCA)", heading_style))
            elements.append(Paragraph(self.data.get('nca_clause'), normal_style))
            section_num += 1
        
        if self.data.get('renewal_clause'):
            elements.append(Paragraph(f"{section_num}. RENEWAL TERMS", heading_style))
            elements.append(Paragraph(self.data.get('renewal_clause'), normal_style))
            section_num += 1
        
        if self.data.get('conveyance_clause'):
            elements.append(Paragraph(f"{section_num}. CONVEYANCE", heading_style))
            elements.append(Paragraph(self.data.get('conveyance_clause'), normal_style))
            section_num += 1
        
        # Scope of Work
        elements.append(Paragraph(f"{section_num}. SCOPE OF WORK (SOW)", heading_style))
        section_num += 1
        
        sow_table_data = self.data.get('sow_table', [])
        if sow_table_data:
            sow_data = [['#', 'Category', 'Title', 'Description', 'Timeline']]
            for idx, item in enumerate(sow_table_data, 1):
                desc = item.get('description', '')
                deliverables = item.get('deliverables', [])
                if deliverables:
                    desc += "\nDeliverables:\n" + "\n".join([f"• {d}" for d in deliverables[:3]])  # Limit for space
                
                timeline = item.get('timeline_weeks')
                timeline_str = f"{timeline}w" if timeline else "TBD"
                
                sow_data.append([
                    str(idx),
                    item.get('category', ''),
                    item.get('title', '')[:30],  # Truncate for table
                    desc[:100] + "..." if len(desc) > 100 else desc,
                    timeline_str
                ])
            
            sow_table = Table(sow_data, colWidths=[0.5*cm, 2.5*cm, 3.5*cm, 7*cm, 1.5*cm])
            sow_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(sow_table)
        else:
            elements.append(Paragraph("(SOW details to be defined)", normal_style))
        elements.append(Spacer(1, 10))
        
        # Project Details
        elements.append(Paragraph(f"{section_num}. PROJECT DETAILS", heading_style))
        section_num += 1
        
        project = self.data.get('project_details', {})
        elements.append(Paragraph(f"<b>Start Date:</b> {format_date(project.get('start_date'))}", normal_style))
        duration = project.get('duration_months')
        elements.append(Paragraph(f"<b>Duration:</b> {duration} months" if duration else "<b>Duration:</b> TBD", normal_style))
        if project.get('payment_schedule'):
            elements.append(Paragraph(f"<b>Payment Schedule:</b> {project.get('payment_schedule')}", normal_style))
        
        # Team Engagement
        elements.append(Paragraph(f"{section_num}. TEAM ENGAGEMENT", heading_style))
        section_num += 1
        
        team_data = self.data.get('team_engagement', [])
        if team_data and isinstance(team_data, list):
            team_table_data = [['Type', 'Count', 'Meetings', 'Hours', 'Rate']]
            for item in team_data:
                team_table_data.append([
                    item.get('type', ''),
                    str(item.get('count', 1)),
                    str(item.get('meetings', 0)),
                    str(item.get('hours', 0)),
                    format_currency(item.get('rate', 12500))
                ])
            
            team_table = Table(team_table_data, colWidths=[4*cm, 2*cm, 2.5*cm, 2*cm, 3*cm])
            team_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            elements.append(team_table)
        else:
            elements.append(Paragraph("(Team details as per SOW)", normal_style))
        elements.append(Spacer(1, 10))
        
        # Pricing
        elements.append(Paragraph(f"{section_num}. PRICING PLAN", heading_style))
        section_num += 1
        
        pricing = self.data.get('pricing', {})
        pricing_data = [
            ['Item', 'Amount'],
            ['Total Meetings', str(pricing.get('total_meetings', 0))],
            ['Subtotal', format_currency(pricing.get('subtotal', 0))],
            ['Discount', format_currency(pricing.get('discount', 0))],
            ['GST (18%)', format_currency(pricing.get('gst', 0))],
            ['Grand Total', format_currency(pricing.get('total', 0))],
        ]
        
        pricing_table = Table(pricing_data, colWidths=[6*cm, 6*cm])
        pricing_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        elements.append(pricing_table)
        elements.append(Spacer(1, 10))
        
        # Payment Terms
        elements.append(Paragraph(f"{section_num}. PAYMENT TERMS & CONDITIONS", heading_style))
        section_num += 1
        
        elements.append(Paragraph(f"<b>Payment Terms:</b> {self.data.get('payment_terms', 'Net 30 days')}", normal_style))
        if self.data.get('payment_conditions'):
            elements.append(Paragraph(self.data.get('payment_conditions'), normal_style))
        if self.data.get('special_conditions'):
            elements.append(Paragraph(f"<b>Special Conditions:</b> {self.data.get('special_conditions')}", normal_style))
        
        # Signatures
        elements.append(PageBreak())
        elements.append(Paragraph(f"{section_num}. SIGNATURES", heading_style))
        elements.append(Spacer(1, 30))
        
        sig_data = [
            ['For D&V Business Consulting', 'For Client'],
            ['\n\n\n_____________________', '\n\n\n_____________________'],
            ['Authorized Signatory', party_name or 'Client Representative'],
            ['Date: _______________', 'Date: _______________'],
        ]
        
        sig_table = Table(sig_data, colWidths=[8*cm, 8*cm])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
        ]))
        elements.append(sig_table)
        
        doc.build(elements)
        buffer.seek(0)
        return buffer


class SOWDocumentGenerator:
    """Generates standalone SOW documents in Word and PDF formats"""
    
    COMPANY_NAME = "D&V Business Consulting"
    
    def __init__(self, sow_data: Dict[str, Any], lead_data: Optional[Dict[str, Any]] = None, 
                 pricing_plan_data: Optional[Dict[str, Any]] = None):
        self.sow = sow_data
        self.lead = lead_data or {}
        self.pricing_plan = pricing_plan_data or {}
    
    def generate_word(self) -> BytesIO:
        """Generate Word document for SOW"""
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Title
        title = doc.add_heading('SCOPE OF WORK', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Client Info
        if self.lead:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Client: {self.lead.get('company', '')}")
            doc.add_paragraph()
        
        # SOW Status
        overall_status = self.sow.get('overall_status', 'draft')
        p = doc.add_paragraph()
        p.add_run("Status: ").bold = True
        p.add_run(overall_status.replace('_', ' ').title())
        
        p = doc.add_paragraph()
        p.add_run("Version: ").bold = True
        p.add_run(str(self.sow.get('current_version', 1)))
        
        doc.add_paragraph()
        
        # SOW Items Table
        doc.add_heading('Deliverables', level=1)
        
        items = self.sow.get('items', [])
        if items:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            headers = ['#', 'Category', 'Title', 'Description', 'Consultant', 'Timeline']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for idx, item in enumerate(items, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = item.get('category', '').replace('_', ' ').title()
                row[2].text = item.get('title', '')
                
                # Description with deliverables
                desc = item.get('description', '')
                deliverables = item.get('deliverables', [])
                if deliverables:
                    desc += "\n" + "\n".join([f"• {d}" for d in deliverables])
                row[3].text = desc
                
                row[4].text = item.get('assigned_consultant_name', 'TBD')
                
                timeline = item.get('timeline_weeks')
                start = item.get('start_week')
                timeline_str = f"Week {start}" if start else ""
                if timeline:
                    timeline_str += f" ({timeline}w)"
                row[5].text = timeline_str or "TBD"
        else:
            doc.add_paragraph("No items defined yet.")
        
        doc.add_paragraph()
        
        # Summary
        doc.add_heading('Summary', level=1)
        
        total_items = len(items)
        approved_count = len([i for i in items if i.get('status') == 'approved'])
        completed_count = len([i for i in items if i.get('status') == 'completed'])
        
        p = doc.add_paragraph()
        p.add_run("Total Items: ").bold = True
        p.add_run(str(total_items))
        
        p = doc.add_paragraph()
        p.add_run("Approved: ").bold = True
        p.add_run(str(approved_count))
        
        p = doc.add_paragraph()
        p.add_run("Completed: ").bold = True
        p.add_run(str(completed_count))
        
        # Save
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_pdf(self) -> BytesIO:
        """Generate PDF document for SOW"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], 
                                      fontSize=18, alignment=TA_CENTER, spaceAfter=20)
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], 
                                        fontSize=12, spaceBefore=15, spaceAfter=8)
        normal_style = ParagraphStyle('Normal', parent=styles['Normal'], 
                                       fontSize=10, spaceAfter=6)
        
        elements = []
        
        # Title
        elements.append(Paragraph("SCOPE OF WORK", title_style))
        
        # Client Info
        if self.lead:
            elements.append(Paragraph(f"Client: {self.lead.get('company', '')}", 
                                      ParagraphStyle('Center', parent=normal_style, alignment=TA_CENTER)))
        elements.append(Spacer(1, 20))
        
        # Status
        overall_status = self.sow.get('overall_status', 'draft')
        elements.append(Paragraph(f"<b>Status:</b> {overall_status.replace('_', ' ').title()}", normal_style))
        elements.append(Paragraph(f"<b>Version:</b> {self.sow.get('current_version', 1)}", normal_style))
        elements.append(Spacer(1, 15))
        
        # Items Table
        elements.append(Paragraph("Deliverables", heading_style))
        
        items = self.sow.get('items', [])
        if items:
            table_data = [['#', 'Category', 'Title', 'Consultant', 'Timeline', 'Status']]
            
            for idx, item in enumerate(items, 1):
                timeline = item.get('timeline_weeks')
                start = item.get('start_week')
                timeline_str = f"W{start}" if start else ""
                if timeline:
                    timeline_str += f" ({timeline}w)"
                
                status = item.get('status', 'draft').replace('_', ' ').title()
                
                table_data.append([
                    str(idx),
                    item.get('category', '').replace('_', ' ').title()[:15],
                    item.get('title', '')[:25],
                    item.get('assigned_consultant_name', 'TBD')[:15],
                    timeline_str or "TBD",
                    status
                ])
            
            table = Table(table_data, colWidths=[0.7*cm, 2.5*cm, 4*cm, 3*cm, 2*cm, 2.5*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(table)
        else:
            elements.append(Paragraph("No items defined yet.", normal_style))
        
        elements.append(Spacer(1, 20))
        
        # Summary
        elements.append(Paragraph("Summary", heading_style))
        
        total_items = len(items)
        approved_count = len([i for i in items if i.get('status') == 'approved'])
        completed_count = len([i for i in items if i.get('status') == 'completed'])
        
        elements.append(Paragraph(f"<b>Total Items:</b> {total_items}", normal_style))
        elements.append(Paragraph(f"<b>Approved:</b> {approved_count}", normal_style))
        elements.append(Paragraph(f"<b>Completed:</b> {completed_count}", normal_style))
        
        doc.build(elements)
        buffer.seek(0)
        return buffer

"""
HR Module Documentation Generator
Generates comprehensive ERP documentation in PDF and DOCX formats
"""
import os
from datetime import datetime, timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF
import uuid


class HRDocumentationGenerator:
    """Generate HR Module documentation in PDF and DOCX formats"""
    
    def __init__(self):
        self.company_name = "DVBC Consulting"
        self.system_name = "NETRA ERP"
        self.version = "2.0"
        self.generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
        
    def get_hr_documentation_content(self):
        """Get all HR module documentation content"""
        return {
            "title": "NETRA ERP - HR Module Documentation",
            "subtitle": f"Complete HR Operations Guide | Version {self.version}",
            "sections": [
                self._get_system_overview(),
                self._get_business_logic(),
                self._get_roles_permissions(),
                self._get_workflow_maps(),
                self._get_admin_guide(),
                self._get_sops(),
                self._get_training_manual(),
                self._get_troubleshooting(),
                self._get_audit_compliance(),
                self._get_quick_start(),
            ]
        }
    
    def _get_system_overview(self):
        """Section 1: System Overview"""
        return {
            "title": "1. System Overview",
            "content": [
                {
                    "heading": "1.1 Purpose of HR Module",
                    "text": """The HR Module in NETRA ERP is designed to streamline all human resource operations for DVBC Consulting. It provides a centralized platform for managing employees from hiring to retirement, handling attendance tracking, leave management, payroll processing, and compliance documentation.

Key Benefits:
- Paperless employee onboarding
- Automated attendance tracking with policy enforcement
- Streamlined leave approval workflow
- Integrated payroll calculations
- Comprehensive audit trail for all HR actions"""
                },
                {
                    "heading": "1.2 Module Components",
                    "text": """The HR Module consists of the following sub-modules:

1. Employee Management
   - Employee onboarding and profiles
   - Document management
   - Organization hierarchy

2. Attendance Management
   - Daily check-in/check-out
   - Work from home tracking
   - Attendance approvals

3. Leave Management
   - Leave requests and approvals
   - Leave balance tracking
   - Holiday calendar

4. Payroll Management
   - Salary slip generation
   - CTC structure design
   - Deductions and allowances

5. Compliance & Documentation
   - Offer letters and contracts
   - Bank detail management
   - Audit logs"""
                },
                {
                    "heading": "1.3 Integration with Other Modules",
                    "text": """The HR Module integrates seamlessly with other NETRA modules:

- Sales Module: Employee assignment to leads and deals
- Consulting Module: Consultant allocation and project staffing
- Finance Module: Payroll disbursement and expense tracking
- Admin Module: User access and system configuration"""
                }
            ]
        }
    
    def _get_business_logic(self):
        """Section 2: Business Logic Documentation"""
        return {
            "title": "2. Business Logic Documentation",
            "content": [
                {
                    "heading": "2.1 Employee Onboarding Rules",
                    "text": """Employee ID Generation:
- Format: EMP followed by 3-4 digit number (e.g., EMP001, EMP1234)
- Auto-incremented based on existing employees
- Cannot be modified after creation

Required Fields for Onboarding:
- Personal: First Name, Last Name, Email, Phone, Date of Birth, Gender, Address
- Employment: Employee ID, Designation, Department, Employment Type, Joining Date, Reporting Manager

Department-Based Access:
- Employees are assigned to one or more departments
- Each department grants access to specific pages in the system
- Primary department determines the main dashboard view"""
                },
                {
                    "heading": "2.2 Attendance Policy",
                    "text": """Working Hours:
- Non-Consulting Staff: 10:00 AM - 7:00 PM
- Consulting Staff: 10:30 AM - 7:30 PM
- Working Days: Monday to Saturday

Grace Period Rules:
- 30 minutes grace for check-in/check-out
- Maximum 3 grace days per month allowed
- Beyond 3 days: Rs. 100 penalty per violation

Custom Policies:
- HR can set custom timing for specific employees
- Custom policies override role-based defaults
- Effective dates can be specified"""
                },
                {
                    "heading": "2.3 Leave Policy",
                    "text": """Leave Entitlements (Annual):
- Casual Leave: 12 days
- Sick Leave: 6 days
- Earned Leave: 15 days

Leave Request Rules:
- Employee submits request with dates and reason
- Reporting Manager approval required
- Half-day leave supported (First Half / Second Half)
- Balance checked before approval

Leave Balance Calculation:
- Available = Total Entitled - Used
- Insufficient balance blocks request submission"""
                },
                {
                    "heading": "2.4 Bank Detail Change Workflow",
                    "text": """Two-Level Approval Process:

Step 1: Employee Request
- Employee submits new bank details
- Supporting document (cancelled cheque/passbook) required
- IFSC code auto-verified via Razorpay API

Step 2: HR Approval
- HR Manager reviews and verifies details
- Can approve or reject with remarks
- Approved requests move to Admin

Step 3: Admin Final Approval
- Admin gives final sign-off
- On approval, employee bank details updated
- Employee notified via in-app notification"""
                },
                {
                    "heading": "2.5 Probation Rules",
                    "text": """Probation Period:
- Consulting Department: 6 months
- Other Departments: 3 months

Confirmation Date:
- Auto-calculated: Joining Date + Probation Period
- Can be extended by HR with reason"""
                }
            ]
        }
    
    def _get_roles_permissions(self):
        """Section 3: Role-Based Access & Permissions Matrix"""
        return {
            "title": "3. Role-Based Access & Permissions",
            "content": [
                {
                    "heading": "3.1 HR Roles Overview",
                    "text": """HR Executive:
- View employee records
- Record attendance
- View leave requests
- Upload documents

HR Manager:
- All HR Executive permissions
- Approve/reject leave requests
- Approve bank change requests (first level)
- Credit leave balances
- Set custom attendance policies
- Generate payroll reports

Admin:
- All HR Manager permissions
- Final approval for bank changes
- Delete employee records
- System configuration"""
                },
                {
                    "heading": "3.2 Permission Matrix",
                    "text": """| Action                          | Employee | HR Executive | HR Manager | Admin |
|--------------------------------|----------|--------------|------------|-------|
| View Own Profile               | Yes      | Yes          | Yes        | Yes   |
| View All Employees             | No       | Yes          | Yes        | Yes   |
| Create Employee                | No       | Yes          | Yes        | Yes   |
| Edit Employee Details          | No       | Yes          | Yes        | Yes   |
| Delete Employee                | No       | No           | No         | Yes   |
| Record Attendance              | Own      | All          | All        | All   |
| Approve Attendance             | No       | Yes          | Yes        | Yes   |
| Submit Leave Request           | Yes      | Yes          | Yes        | Yes   |
| Approve Leave Request          | No       | No           | Yes        | Yes   |
| View Payroll                   | Own      | No           | Yes        | Yes   |
| Process Payroll                | No       | No           | Yes        | Yes   |
| Bank Change Request            | Submit   | View         | Approve L1 | Approve L2 |
| Custom Attendance Policy       | No       | No           | Yes        | Yes   |"""
                },
                {
                    "heading": "3.3 Page Access by Department",
                    "text": """HR Department Pages:
/employees - Employee directory
/onboarding - New employee onboarding
/attendance - Attendance management
/leave-mgmt - Leave management
/payroll - Payroll processing
/org-chart - Organization hierarchy
/letter-management - Offer letters and contracts
/letterhead-settings - Letter templates
/ctc-designer - CTC structure design
/attendance-approvals - Attendance approvals

My Workspace (All Employees):
/my-attendance - Personal attendance
/my-leaves - Personal leave requests
/my-salary-slips - Personal salary slips
/my-expenses - Personal expenses"""
                }
            ]
        }
    
    def _get_workflow_maps(self):
        """Section 4: End-to-End Workflow Maps"""
        return {
            "title": "4. End-to-End Workflow Maps",
            "content": [
                {
                    "heading": "4.1 Employee Onboarding Workflow",
                    "text": """Step 1: Quick Import (Optional)
- Download CSV template
- Fill employee details
- Upload for bulk onboarding
OR proceed with single employee

Step 2: Personal Information
- Enter name, email, phone
- Date of birth, gender
- Address and emergency contact

Step 3: Employment Details
- Employee ID (auto-generated)
- Designation, Department selection
- Employment type, Joining date
- Reporting manager assignment

Step 4: Document Upload
- Photo, ID proof
- Education certificates
- Experience letters

Step 5: Bank Details
- Account number, IFSC code
- Bank name and branch
- Cancelled cheque upload

Step 6: Review & Submit
- Verify all information
- System creates employee record
- Portal access granted automatically
- Welcome email sent (if configured)"""
                },
                {
                    "heading": "4.2 Leave Request Workflow",
                    "text": """Step 1: Employee Submits Request
- Select leave type (Casual/Sick/Earned)
- Choose dates (start and end)
- Option for half-day leave
- Add reason for leave

Step 2: Balance Check
- System validates available balance
- Blocks if insufficient

Step 3: Reporting Manager Review
- RM receives notification
- Reviews request details
- Approves or Rejects

Step 4: Balance Update (On Approval)
- Leave balance deducted
- Calendar updated
- Employee notified

Step 5: HR Notification
- HR receives informational update
- No action required"""
                },
                {
                    "heading": "4.3 Daily Attendance Workflow",
                    "text": """Morning: Check-In
- Employee opens Quick Check-In
- System captures timestamp
- Location recorded (if enabled)

Evening: Check-Out
- Employee completes check-out
- Working hours calculated
- Status determined (Present/Half Day)

Monthly: Validation
- HR runs auto-validation
- System applies attendance policy
- Grace violations identified
- Penalties calculated

Payroll: Integration
- Validated attendance linked to payroll
- Deductions applied automatically
- Salary slip generated"""
                },
                {
                    "heading": "4.4 Bank Change Request Workflow",
                    "text": """Step 1: Employee Initiates
- Navigate to Profile > Bank Details
- Enter new bank information
- Upload proof document
- Submit request

Step 2: HR Manager Review
- Receives notification
- Verifies documents
- Approves (moves to Admin) or Rejects

Step 3: Admin Final Approval
- Reviews HR-approved request
- Gives final approval
- Bank details updated in system

Step 4: Confirmation
- Employee receives notification
- New details active for next payroll"""
                }
            ]
        }
    
    def _get_admin_guide(self):
        """Section 5: Configuration Guide (Admin Manual)"""
        return {
            "title": "5. Configuration Guide (Admin Manual)",
            "content": [
                {
                    "heading": "5.1 Managing Users",
                    "text": """Creating New Users:
1. Navigate to Admin > Users
2. Click 'Add User'
3. Fill required details:
   - Full Name
   - Email (must be unique)
   - Role selection
   - Department assignment
4. Save - temporary password generated

Resetting Passwords:
1. Find user in Users list
2. Click 'Reset Password'
3. New temporary password generated
4. User must change on first login

Deactivating Users:
1. Find user in list
2. Toggle 'Active' status to Off
3. User loses system access immediately"""
                },
                {
                    "heading": "5.2 Attendance Policy Configuration",
                    "text": """Default Policy Settings:
- Located in: Settings > Attendance Policy
- Working Days: Select applicable days
- Office Hours: Set check-in/out times
- Grace Period: Configure grace minutes
- Penalty Amount: Set per-violation fee

Custom Employee Policies:
1. Go to Attendance > Custom Policies
2. Search for employee
3. Set custom timing
4. Add reason for exception
5. Set effective dates
6. Save policy"""
                },
                {
                    "heading": "5.3 Leave Configuration",
                    "text": """Setting Leave Entitlements:
1. Go to Settings > Leave Policy
2. Configure per leave type:
   - Casual Leave days
   - Sick Leave days
   - Earned Leave days
3. Set carry-forward rules
4. Configure accrual frequency

Holiday Calendar:
1. Navigate to Settings > Holidays
2. Add public holidays
3. Specify optional/restricted holidays
4. Holidays appear in leave calendar"""
                },
                {
                    "heading": "5.4 Letter Templates",
                    "text": """Managing Offer Letter Templates:
1. Go to Letter Management > Templates
2. Create or edit template
3. Use placeholders:
   - {{employee_name}}
   - {{designation}}
   - {{department}}
   - {{salary}}
   - {{joining_date}}
4. Preview and save

Letterhead Settings:
1. Navigate to Letterhead Settings
2. Upload company logo
3. Set header/footer text
4. Configure signature fields"""
                }
            ]
        }
    
    def _get_sops(self):
        """Section 6: Standard Operating Procedures"""
        return {
            "title": "6. Standard Operating Procedures (SOPs)",
            "content": [
                {
                    "heading": "6.1 Daily Tasks - HR Executive",
                    "text": """Morning (9:30 AM - 10:30 AM):
1. Check attendance dashboard
2. Note any missing check-ins
3. Review pending approvals

Mid-day (2:00 PM - 3:00 PM):
1. Process document verifications
2. Update employee records as needed
3. Respond to employee queries

End of Day (6:00 PM - 7:00 PM):
1. Run attendance summary report
2. Flag incomplete attendance
3. Update pending tasks list"""
                },
                {
                    "heading": "6.2 Weekly Tasks - HR Manager",
                    "text": """Monday:
- Review attendance summary (previous week)
- Approve pending leave requests
- Check bank change requests

Wednesday:
- Review new employee onboarding status
- Conduct policy exception reviews
- Team status meeting

Friday:
- Generate weekly HR reports
- Plan next week activities
- Backup critical data"""
                },
                {
                    "heading": "6.3 Monthly Tasks",
                    "text": """1st Week:
- Run attendance validation
- Apply late penalties
- Process previous month payroll

2nd Week:
- Generate salary slips
- Review leave balances
- Update policy configurations

3rd Week:
- Conduct compliance check
- Review audit logs
- Address escalations

4th Week:
- Prepare monthly HR report
- Plan next month activities
- Backup and archive records"""
                },
                {
                    "heading": "6.4 New Employee Onboarding SOP",
                    "text": """Day -7 (Before Joining):
1. Prepare workstation
2. Create system accounts
3. Prepare welcome kit

Day 1 (Joining Day):
1. Complete onboarding form in NETRA
2. Collect and verify documents
3. Grant portal access
4. Introduce to team

Week 1:
1. Complete department orientation
2. Assign reporting manager
3. Set up attendance tracking
4. Configure leave balance

Day 30:
1. Conduct 30-day review
2. Address any concerns
3. Update records as needed"""
                }
            ]
        }
    
    def _get_training_manual(self):
        """Section 7: Training Manual"""
        return {
            "title": "7. Training Manual",
            "content": [
                {
                    "heading": "7.1 For New HR Staff",
                    "text": """Module 1: System Navigation (Day 1)
- Login and dashboard overview
- Understanding the sidebar menu
- Profile and settings
- Practice: Navigate all HR pages

Module 2: Employee Management (Day 2)
- Adding new employees
- Searching and filtering
- Editing employee records
- Practice: Onboard a test employee

Module 3: Attendance & Leave (Day 3)
- Recording attendance
- Processing leave requests
- Understanding policies
- Practice: Process 5 leave requests

Module 4: Payroll Basics (Day 4)
- Viewing payroll inputs
- Understanding CTC structure
- Generating salary slips
- Practice: Review payroll for 3 employees

Module 5: Reports & Compliance (Day 5)
- Generating HR reports
- Understanding audit logs
- Best practices
- Assessment quiz"""
                },
                {
                    "heading": "7.2 For Employees (Self-Service)",
                    "text": """Accessing Your Dashboard:
1. Login at: your-company-url/login
2. Use Employee ID and password
3. Change password on first login

Marking Attendance:
1. Click 'Quick Attendance' card
2. Allow location (if prompted)
3. Click 'Check In'
4. Click 'Check Out' at day end

Applying for Leave:
1. Go to My Workspace > My Leaves
2. Click 'Apply Leave'
3. Select leave type
4. Choose dates
5. Add reason
6. Submit

Viewing Salary Slips:
1. Go to My Workspace > My Salary Slips
2. Select month
3. View or download slip"""
                },
                {
                    "heading": "7.3 Frequently Asked Questions",
                    "text": """Q: How do I reset my password?
A: Click 'Forgot Password' on login page, or contact HR.

Q: Why can't I apply for leave?
A: Check your leave balance. Insufficient balance prevents submission.

Q: My attendance is showing absent but I was present?
A: Contact HR to verify and correct the record.

Q: How do I change my bank details?
A: Go to Profile > Bank Details > Request Change. HR approval required.

Q: When are salary slips available?
A: Typically by 5th of following month after payroll processing.

Q: How do I view my leave balance?
A: Go to My Workspace > My Leaves. Balance shown at top.

Q: Can I apply for half-day leave?
A: Yes, select 'Half Day' option and choose First or Second half."""
                }
            ]
        }
    
    def _get_troubleshooting(self):
        """Section 8: Troubleshooting Guide"""
        return {
            "title": "8. Troubleshooting Guide",
            "content": [
                {
                    "heading": "8.1 Login Issues",
                    "text": """Problem: Cannot login
Solutions:
1. Verify Employee ID is correct (case-sensitive)
2. Check if account is active
3. Try password reset
4. Clear browser cache
5. Contact HR if issue persists

Problem: Session expires frequently
Solutions:
1. Check internet connection
2. Disable browser extensions
3. Use recommended browser (Chrome/Firefox)
4. Contact IT if continues"""
                },
                {
                    "heading": "8.2 Attendance Issues",
                    "text": """Problem: Check-in button not working
Solutions:
1. Refresh the page
2. Check internet connection
3. Try different browser
4. Contact HR to manually record

Problem: Location not detected
Solutions:
1. Allow location permission in browser
2. Check GPS is enabled
3. Try refreshing page
4. Attendance can be recorded without location

Problem: Wrong attendance status
Solutions:
1. Contact HR immediately
2. Provide correct check-in/out times
3. HR can manually correct records"""
                },
                {
                    "heading": "8.3 Leave Request Issues",
                    "text": """Problem: Insufficient balance error
Solutions:
1. Check leave balance in My Leaves
2. Verify leave type selected
3. Try applying for fewer days
4. Contact HR for adjustment if incorrect

Problem: Request stuck in pending
Solutions:
1. Check if RM is assigned
2. Remind RM to approve
3. Contact HR if RM unavailable

Problem: Cannot cancel approved leave
Solutions:
1. Contact HR to cancel
2. HR can adjust leave balance
3. New request may be needed"""
                },
                {
                    "heading": "8.4 Payroll Issues",
                    "text": """Problem: Salary slip not visible
Solutions:
1. Check if payroll processed for month
2. Verify with HR
3. May take 2-3 days after month end

Problem: Incorrect salary amount
Solutions:
1. Review salary components
2. Check attendance deductions
3. Verify CTC structure
4. Raise ticket with HR

Problem: Bank details incorrect on slip
Solutions:
1. Initiate bank change request
2. Next payroll will use new details
3. Current payment uses old details"""
                }
            ]
        }
    
    def _get_audit_compliance(self):
        """Section 9: Audit & Compliance Controls"""
        return {
            "title": "9. Audit & Compliance Controls",
            "content": [
                {
                    "heading": "9.1 What is Tracked",
                    "text": """Employee Record Changes:
- All edits to employee data logged
- Who made the change
- When it was made
- Before and after values

Attendance Records:
- Check-in/check-out timestamps
- Location data (if enabled)
- Manual corrections with approver

Leave Transactions:
- Request submissions
- Approvals and rejections
- Balance adjustments

Payroll Actions:
- Salary slip generation
- Deductions applied
- Disbursement records

Access Logs:
- Login attempts (success/failure)
- Page access history
- Sensitive data views"""
                },
                {
                    "heading": "9.2 Compliance Reports",
                    "text": """Available Reports:
1. Attendance Compliance Report
   - Attendance rate by department
   - Late arrival patterns
   - Penalty summary

2. Leave Utilization Report
   - Leave balance status
   - Leave patterns
   - Unused leave liability

3. Employee Change Log
   - All modifications
   - Approver details
   - Audit timeline

4. Payroll Audit Trail
   - Salary components
   - Deductions applied
   - Payment confirmations"""
                },
                {
                    "heading": "9.3 Data Retention",
                    "text": """Retention Periods:
- Employee Records: 7 years after exit
- Attendance Data: 3 years
- Payroll Records: 8 years
- Audit Logs: 5 years

Backup Schedule:
- Daily incremental backup
- Weekly full backup
- Monthly archive to cold storage

Data Access:
- Role-based access control
- Sensitive data encrypted
- Access logged for audit"""
                }
            ]
        }
    
    def _get_quick_start(self):
        """Section 10: Quick Start Guide"""
        return {
            "title": "10. Quick Start Guide - Start in 30 Minutes",
            "content": [
                {
                    "heading": "10.1 For HR Manager (First-Time Setup)",
                    "text": """Minutes 1-5: Login & Explore
1. Login with admin-provided credentials
2. Change your password
3. Explore HR Dashboard
4. Note key metrics displayed

Minutes 6-15: Onboard Your First Employee
1. Go to Onboarding page
2. Fill personal information
3. Add employment details
4. Skip documents for now
5. Grant portal access

Minutes 16-20: Configure Basics
1. Review attendance policy
2. Check leave entitlements
3. Note any needed changes

Minutes 21-25: Test Key Workflows
1. Record attendance for test employee
2. Submit a test leave request
3. Approve the request

Minutes 26-30: Review & Plan
1. Check dashboard updated
2. Note questions for IT/Admin
3. Plan full team onboarding"""
                },
                {
                    "heading": "10.2 For New Employee (First Day)",
                    "text": """Minutes 1-5: Access Your Account
1. Receive login credentials from HR
2. Login at: your-company-url
3. Change password immediately
4. Note your Employee ID

Minutes 6-10: Complete Profile
1. Go to My Profile
2. Verify personal information
3. Update emergency contact
4. Review reporting manager

Minutes 11-15: Mark First Attendance
1. Find Quick Attendance on dashboard
2. Click Check In
3. Note your check-in time
4. Remember to check out later

Minutes 16-20: Explore Features
1. View My Leaves (check balance)
2. See My Attendance history
3. Access My Salary Slips (when available)

Minutes 21-25: Apply Test Leave
1. Go to My Leaves
2. Click Apply Leave
3. Select Casual Leave
4. Choose tomorrow
5. Submit request

Minutes 26-30: Familiarize
1. Browse help section
2. Note HR contact
3. Bookmark important pages"""
                },
                {
                    "heading": "10.3 Key Contacts",
                    "text": """For HR Issues:
- Email: hr@dvconsulting.co.in
- Phone: Contact your HR representative

For System Issues:
- Email: support@dvconsulting.co.in
- Raise ticket through IT portal

For Urgent Matters:
- Contact your Reporting Manager
- Escalate to HR Manager

Portal URL: https://your-netra-url.com
Working Hours Support: 10 AM - 7 PM (Mon-Sat)"""
                }
            ]
        }

    def generate_docx(self, output_path: str):
        """Generate DOCX documentation"""
        doc = Document()
        content = self.get_hr_documentation_content()
        
        # Title
        title = doc.add_heading(content["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(content["subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meta info
        meta = doc.add_paragraph(f"Generated: {self.generated_date} | {self.company_name}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Table of Contents header
        doc.add_heading("Table of Contents", level=1)
        for section in content["sections"]:
            doc.add_paragraph(section["title"], style='List Number')
        
        doc.add_page_break()
        
        # Add all sections
        for section in content["sections"]:
            doc.add_heading(section["title"], level=1)
            
            for item in section["content"]:
                doc.add_heading(item["heading"], level=2)
                
                # Handle text with potential tables
                text = item["text"]
                paragraphs = text.strip().split('\n\n')
                
                for para in paragraphs:
                    if para.strip().startswith('|'):
                        # This is a markdown table - add as formatted text
                        table_para = doc.add_paragraph()
                        table_para.add_run(para).font.size = Pt(9)
                    else:
                        doc.add_paragraph(para.strip())
            
            doc.add_page_break()
        
        doc.save(output_path)
        return output_path
    
    def generate_pdf(self, output_path: str):
        """Generate PDF documentation"""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        content = self.get_hr_documentation_content()
        
        # Title page
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 24)
        pdf.cell(0, 60, '', 0, 1)  # Spacing
        pdf.multi_cell(0, 10, content["title"], 0, 'C')
        
        pdf.set_font('Helvetica', '', 14)
        pdf.cell(0, 20, '', 0, 1)
        pdf.multi_cell(0, 8, content["subtitle"], 0, 'C')
        
        pdf.cell(0, 20, '', 0, 1)
        pdf.set_font('Helvetica', 'I', 10)
        pdf.multi_cell(0, 6, f"Generated: {self.generated_date}", 0, 'C')
        pdf.multi_cell(0, 6, self.company_name, 0, 'C')
        
        # Table of Contents
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 18)
        pdf.cell(0, 10, 'Table of Contents', 0, 1)
        pdf.cell(0, 5, '', 0, 1)
        
        pdf.set_font('Helvetica', '', 12)
        for i, section in enumerate(content["sections"], 1):
            pdf.cell(0, 8, f"{section['title']}", 0, 1)
        
        # Add sections
        for section in content["sections"]:
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 16)
            pdf.multi_cell(0, 10, section["title"])
            pdf.cell(0, 5, '', 0, 1)
            
            for item in section["content"]:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.multi_cell(0, 8, item["heading"])
                
                pdf.set_font('Helvetica', '', 10)
                # Clean and encode text properly
                text = item["text"].strip()
                # Replace special characters
                text = text.replace('\u2019', "'").replace('\u2018', "'")
                text = text.replace('\u201c', '"').replace('\u201d', '"')
                text = text.replace('\u2013', '-').replace('\u2014', '-')
                text = text.encode('latin-1', 'replace').decode('latin-1')
                
                pdf.multi_cell(0, 6, text)
                pdf.cell(0, 5, '', 0, 1)
        
        pdf.output(output_path)
        return output_path


def generate_hr_documentation(output_dir: str = "/tmp"):
    """Generate both PDF and DOCX documentation files"""
    generator = HRDocumentationGenerator()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    docx_path = os.path.join(output_dir, f"NETRA_HR_Module_Documentation_{timestamp}.docx")
    pdf_path = os.path.join(output_dir, f"NETRA_HR_Module_Documentation_{timestamp}.pdf")
    
    generator.generate_docx(docx_path)
    generator.generate_pdf(pdf_path)
    
    return {
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "generated_at": timestamp
    }
